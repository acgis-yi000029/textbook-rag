#!/usr/bin/env python3
"""
Universal Markdown to DOCX Converter for Lab Reports

Usage:
    python convert_md_to_docx.py <input.md> [output.docx]
    
Example:
    python convert_md_to_docx.py Lab1_Template.md Lab1.docx

Features:
    - Auto-preprocesses markdown to remove image alt text (prevents captions in Word)
    - Auto-removes horizontal rules (---) that appear as lines in Word
    - Handles relative image paths
    - Auto-installs pandoc if needed
"""
import sys
import os
import re
import tempfile
import pypandoc
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def preprocess_markdown(md_content):
    """
    Preprocess markdown content before conversion.

    - Preserves YAML frontmatter (--- delimited block at start of file)
    - Removes horizontal rules (---) in the body that appear as lines in Word
    - Removes image alt text to prevent it from appearing as captions in Word
    - Example: ![Step 6 Code](path.png) -> ![](path.png)
    """
    # 1. Separate YAML frontmatter from body content
    # YAML frontmatter starts with --- on line 1 and ends with --- on another line
    frontmatter = ''
    body = md_content
    
    if md_content.startswith('---'):
        # Find the closing --- of frontmatter
        end_match = re.search(r'\n---\s*\n', md_content[3:])
        if end_match:
            split_pos = 3 + end_match.end()
            frontmatter = md_content[:split_pos]
            body = md_content[split_pos:]
    
    # 2. Remove horizontal rules (--- on its own line) from BODY only
    # These appear as visible lines in Word documents
    body = re.sub(r'^---\s*$\n?', '', body, flags=re.MULTILINE)

    # 3. Remove image alt text
    # Pattern matches ![any text](path) and replaces with ![](path)
    # This prevents alt text from showing as captions in Word documents
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'

    def replace_alt_text(match):
        image_path = match.group(2)
        return f'![]({image_path})'

    body = re.sub(pattern, replace_alt_text, body)
    return frontmatter + body


def ensure_pandoc():
    """Ensure pandoc is available, download if necessary."""
    try:
        version = pypandoc.get_pandoc_version()
        print(f"[OK] Pandoc version {version} found")
        return True
    except OSError:
        print("[WARN] Pandoc not found. Downloading...")
        try:
            pypandoc.download_pandoc()
            print("[OK] Pandoc downloaded successfully!")
            return True
        except Exception as e:
            print(f"[ERROR] Failed to download pandoc: {e}")
            print("\nPlease install pandoc manually:")
            print("  Windows: choco install pandoc")
            print("  Or download from: https://pandoc.org/installing.html")
            return False


def _add_page_break_after(paragraph):
    """Add a page break after a paragraph by setting pageBreakBefore on the NEXT paragraph."""
    # Add a page break at the end of this paragraph
    p = paragraph._element
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    
    run = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run.append(br)
    p.append(run)


def _format_cover_page(doc):
    """
    Format cover page: center vertically, adjust font sizes.

    Cover page elements (by style):
    - Title: main title, 24pt bold
    - Subtitle: course info, 14pt
    - Author: student name/ID, 12pt
    - Date: submission date, 12pt
    """
    title_styles = {'Title', 'Subtitle', 'Author', 'Date'}
    cover_paragraphs = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name in title_styles:
            cover_paragraphs.append((style_name, para))
        else:
            break  # Stop at first non-cover paragraph

    if not cover_paragraphs:
        return

    # Format each cover element
    for style_name, para in cover_paragraphs:
        # Center align
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if style_name == 'Title':
            # Add vertical spacing to center content (~200pt from top)
            para.paragraph_format.space_before = Pt(200)
            para.paragraph_format.space_after = Pt(24)
            for run in para.runs:
                run.font.size = Pt(24)
                run.font.bold = True
        elif style_name == 'Subtitle':
            para.paragraph_format.space_after = Pt(48)
            for run in para.runs:
                run.font.size = Pt(14)
        elif style_name in ('Author', 'Date'):
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(12)

    print("  [OK] Cover page formatted (centered, font sizes adjusted)")


def _insert_toc_field(doc, insert_before_idx):
    """
    Insert Word native TOC field before a paragraph.
    Only includes Heading 1 entries (Steps).
    """
    para = doc.paragraphs[insert_before_idx]

    # Insert TOC heading
    toc_title = para.insert_paragraph_before('Table of Contents')
    toc_title.style = 'TOC Heading'
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.font.size = Pt(18)
        run.font.bold = True

    # Insert TOC field paragraph
    para = doc.paragraphs[insert_before_idx + 1]
    toc_para = para.insert_paragraph_before('')

    # Create TOC field - only level 1 headings
    run = toc_para.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r' TOC \o "1-1" \h \z '  # Only level 1

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Placeholder
    placeholder_run = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = 'Update field (Ctrl+A, F9) to see TOC'
    placeholder_run.append(placeholder_text)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder_run)
    run._r.append(fldChar3)

    print("  [OK] Word TOC field inserted (level 1 only)")


def _insert_page_breaks(docx_file):
    """
    Post-process docx:
    1. Format cover page (center, font sizes)
    2. Insert Word native TOC field (level 1 only)
    3. Add page breaks after cover page and TOC
    """
    doc = Document(docx_file)
    paragraphs = doc.paragraphs

    # Format cover page first
    _format_cover_page(doc)

    title_styles = {'Title', 'Subtitle', 'Author', 'Date'}

    last_title_idx = -1
    first_heading_idx = -1

    for i, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ''

        if style_name in title_styles:
            last_title_idx = i
        elif style_name == 'Heading 1' and first_heading_idx < 0:
            first_heading_idx = i

    # Insert TOC after cover page, before first heading
    if first_heading_idx > 0:
        _insert_toc_field(doc, first_heading_idx)
        # Indices shifted by 2 (TOC heading + TOC field)
        first_heading_idx += 2
        last_title_idx = last_title_idx  # Cover page idx unchanged

    # Re-read paragraphs after insertion
    paragraphs = doc.paragraphs

    # Add page break after cover page
    if last_title_idx >= 0:
        _add_page_break_after(paragraphs[last_title_idx])
        print(f"  [OK] Page break after cover page (paragraph {last_title_idx})")

    # Add page break before first body heading (after TOC)
    # Find TOC field paragraph (it's 2 paragraphs after cover page now)
    toc_field_idx = last_title_idx + 2 if last_title_idx >= 0 else -1
    if toc_field_idx >= 0 and toc_field_idx < len(paragraphs):
        _add_page_break_after(paragraphs[toc_field_idx])
        print(f"  [OK] Page break after TOC (paragraph {toc_field_idx})")

    doc.save(docx_file)


def convert_md_to_docx(md_file, docx_file=None, reference_doc=None):
    """
    Convert markdown file to docx with proper formatting.
    
    Args:
        md_file: Path to input markdown file
        docx_file: Path to output docx file (optional, defaults to same name)
        reference_doc: Path to reference docx template (optional)
    """
    # Validate input file
    if not os.path.exists(md_file):
        print(f"[ERROR] Error: Input file not found: {md_file}")
        return False
    
    # Determine output file
    if docx_file is None:
        docx_file = os.path.splitext(md_file)[0] + '.docx'
    
    # Get absolute paths
    md_file = os.path.abspath(md_file)
    docx_file = os.path.abspath(docx_file)
    md_dir = os.path.dirname(md_file)
    
    print(f"\n{'='*60}")
    print(f"Converting Markdown to DOCX")
    print(f"{'='*60}")
    print(f"Input:  {md_file}")
    print(f"Output: {docx_file}")
    print(f"{'='*60}\n")
    
    # Prepare pandoc arguments
    # Note: Don't use --toc here, we add Word native TOC field in post-processing
    extra_args = [
        '--standalone',
        f'--resource-path={md_dir}',  # Look for images relative to md file
        '--wrap=preserve',  # Preserve line breaks
    ]
    
    if reference_doc and os.path.exists(reference_doc):
        extra_args.extend(['--reference-doc', reference_doc])
        print(f"Using reference template: {reference_doc}")
    
    try:
        # Read and preprocess markdown content
        print("Preprocessing markdown (removing alt text and horizontal rules)...")
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        processed_content = preprocess_markdown(md_content)
        
        # Create a temporary file with preprocessed content
        # Keep it in the same directory to preserve relative image paths
        temp_md_file = os.path.join(md_dir, '_temp_preprocessed.md')
        with open(temp_md_file, 'w', encoding='utf-8') as f:
            f.write(processed_content)
        
        try:
            # Convert the preprocessed file
            pypandoc.convert_file(
                temp_md_file,
                'docx',
                outputfile=docx_file,
                extra_args=extra_args
            )
        finally:
            # Clean up temporary file
            if os.path.exists(temp_md_file):
                os.remove(temp_md_file)
        
        # Post-process: insert page breaks for cover page and TOC
        print("Post-processing: inserting page breaks...")
        _insert_page_breaks(docx_file)
        
        # Success
        file_size = os.path.getsize(docx_file) / 1024
        print(f"\n[OK] Conversion successful!")
        print(f"  Output file: {docx_file}")
        print(f"  File size: {file_size:.2f} KB")
        
        # Validation checklist
        print(f"\n{'='*60}")
        print("Validation Checklist:")
        print("  [ ] Open the .docx file and verify:")
        print("  [ ] All images display correctly")
        print("  [ ] Headings are properly formatted")
        print("  [ ] Tables are formatted correctly")
        print("  [ ] Page layout is appropriate")
        print(f"{'='*60}\n")
        
        return True
        
    except Exception as e:
        print(f"\n[ERROR] Conversion failed: {e}")
        print("\nTroubleshooting:")
        print("  1. Check that all images exist in the images/ directory")
        print("  2. Verify markdown syntax is correct")
        print("  3. Ensure image paths are relative (e.g., images/pic.png)")
        print("  4. Check for special characters in filenames")
        return False


def main():
    """Main entry point."""
    # Check arguments
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py <input.md> [output.docx]")
        print("\nExample:")
        print("  python convert_md_to_docx.py Lab1_Template.md Lab1.docx")
        sys.exit(1)
    
    md_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # Ensure pandoc is available
    if not ensure_pandoc():
        sys.exit(1)
    
    # Convert
    success = convert_md_to_docx(md_file, docx_file)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
