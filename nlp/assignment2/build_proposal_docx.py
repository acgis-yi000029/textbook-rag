"""
Build a nicely formatted proposal .docx (v2 - with MinerU/DocLayout-YOLO).
生成格式工整的 proposal Word 文档（v2 - 含 MinerU/DocLayout-YOLO 方案）
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        set_cell_shading(cell, '2F5496')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], 'D6E4F0')

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── TITLE ──
    title = doc.add_heading('CST8507 Assignment 2 — Proposal', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'AI Textbook Q&A System:\n'
        'A RAG-Based Educational Question Answering System\n'
        'with Deep Source Tracing'
    )
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(47, 84, 150)

    # ── TEAM ──
    doc.add_heading('Team Details', level=1)
    for name in ['Wang, Peng', 'Yoo, Hye Ran']:
        p = doc.add_paragraph(name, style='List Bullet')
        for r in p.runs:
            r.font.size = Pt(11)

    # ── TOPIC AND MOTIVATION ──
    doc.add_heading('Topic and Motivation', level=1)
    p = doc.add_paragraph()
    p.add_run('We will build an ').font.size = Pt(11)
    r = p.add_run('Educational Question Answering')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        ' system that answers questions about AI/ML concepts using a curated '
        'collection of canonical textbooks as the knowledge base.'
    ).font.size = Pt(11)

    p2 = doc.add_paragraph()
    r = p2.add_run('Motivation: ')
    r.bold = True
    r.font.size = Pt(11)
    p2.add_run(
        'Students studying AI, machine learning, and NLP face the challenge of '
        'navigating across many textbooks to find relevant information. Our system '
        'allows users to ask natural language questions and receive accurate, grounded '
        'answers with deep source tracing — not just citing a book title, but '
        'pinpointing the exact page and region where the answer was found. Users can '
        'click on any source reference to see the original PDF page with the relevant '
        'area highlighted. This promotes transparency and efficient learning.'
    ).font.size = Pt(11)

    # ── DATASET ──
    doc.add_heading('Dataset', level=1)
    p = doc.add_paragraph(
        'Our knowledge base consists of 30+ canonical textbooks in PDF format, covering:'
    )
    p.runs[0].font.size = Pt(11)

    add_table(doc,
        headers=['Domain', 'Books', 'Examples'],
        rows=[
            ['Machine Learning', '7 books', 'ISLR, ESL, PRML, Deep Learning (Goodfellow), PML (Murphy)'],
            ['Mathematics', '5 books', 'MML (Deisenroth), Convex Optimization (Boyd), Information Theory (MacKay)'],
            ['NLP', '3 books', 'SLP3 (Jurafsky), Intro to IR (Manning), NLP Notes (Eisenstein)'],
            ['Reinforcement Learning', '1 book', 'RL: An Introduction (Sutton & Barto)'],
            ['Computer Vision', '1 book', 'Computer Vision (Szeliski)'],
            ['Python', '3 books', 'Fluent Python (Ramalho), Python Cookbook (Beazley), Think Python (Downey)'],
        ],
        col_widths=[4.5, 2.5, 10.5]
    )

    doc.add_heading('Preprocessing Pipeline', level=2)
    steps = [
        ('Layout analysis', ' using MinerU (magic-pdf) with DocLayout-YOLO (GPU-accelerated) to detect document regions (text, tables, formulas, figures) with bounding box coordinates'),
        ('Specialized extraction', ' per content type: text \u2192 plain text, tables \u2192 HTML, formulas \u2192 LaTeX, figures \u2192 image + caption'),
        ('Intelligent chunking', ' with metadata (book title, chapter, page number, bounding box coordinates)'),
        ('Dual indexing:', ' SQLite with FTS5 full-text search index (BM25) + ChromaDB vector store with sentence-transformers embeddings (all-MiniLM-L6-v2)'),
        ('PageIndex tree generation', ' \u2014 hierarchical table-of-contents tree per book for reasoning-based retrieval'),
    ]
    for i, (label, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. ')
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run('Source: ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        'All textbooks are publicly available open-access editions (e.g., from '
        "authors' websites or institutional repositories) or legally obtained "
        'educational copies already in our possession. '
        'Total dataset size: ~500MB of PDF documents.'
    ).font.size = Pt(11)

    # ── APPROACH ──
    doc.add_heading('Approach', level=1)
    p = doc.add_paragraph()
    r = p.add_run('RAG Architecture: ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        'We adopt a hybrid retrieval architecture that combines four complementary '
        'retrieval methods, fused via Reciprocal Rank Fusion (RRF), to maximize '
        'retrieval quality across different query types.'
    ).font.size = Pt(11)

    components = [
        ('Document Processing:', ' MinerU (magic-pdf) for layout-aware PDF parsing \u2014 uses DocLayout-YOLO to detect 10 categories of document elements, preserving bounding box coordinates'),
        ('Chunking:', ' Layout-aware chunking that keeps tables and formulas intact, with source metadata (book, page, section, bbox)'),
        ('Four Retrieval Methods:', ''),
    ]
    for i, (label, desc) in enumerate(components, 1):
        p = doc.add_paragraph(f'{i}. ')
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        if desc:
            p.add_run(desc).font.size = Pt(11)

    # Four retrieval methods as sub-items
    retrieval_methods = [
        ('\u2460 SQLite FTS5 Keyword Search (BM25):', ' Full-text search using SQLite\'s built-in FTS5 extension with BM25 ranking \u2014 fast exact keyword matching'),
        ('\u2461 ChromaDB Vector Search (Semantic):', ' Embedding-based retrieval using sentence-transformers (all-MiniLM-L6-v2) + ChromaDB \u2014 understands synonyms and semantic similarity'),
        ('\u2462 PageIndex Tree Search (LLM Reasoning):', ' Build a hierarchical TOC tree per textbook, then use the LLM to navigate top-down to find relevant sections'),
        ('\u2463 Metadata Filter Search (Structured):', ' Structured filtering by book title, chapter, page number, and content type'),
    ]
    for label, desc in retrieval_methods:
        p = doc.add_paragraph('    ')
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)

    remaining = [
        ('Result Fusion:', ' Reciprocal Rank Fusion (RRF) combines results from all four methods into a single ranked list'),
        ('LLM:', ' Ollama with qwen2.5:0.5b model (~0.4GB memory footprint)'),
        ('UI:', ' Streamlit interface with clickable source references that highlight the exact region on the original PDF page'),
    ]
    for i, (label, desc) in enumerate(remaining, 4):
        p = doc.add_paragraph(f'{i}. ')
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run('Key Feature \u2014 Deep Source Tracing: ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        'Each answer includes references to the exact book, chapter, page, and spatial '
        'location. Clicking a reference renders the original PDF page with the source '
        'region highlighted in a yellow bounding box, enabling users to instantly verify '
        'the answer against the original material.'
    ).font.size = Pt(11)

    # ── PLAN OF WORK ──
    doc.add_heading('Plan of Work', level=1)

    add_table(doc,
        headers=['Week', 'Milestone'],
        rows=[
            ['Week 1 (Mar 3\u20139)', 'MinerU setup, PDF layout analysis, SQLite + ChromaDB schema, batch processing'],
            ['Week 2 (Mar 10\u201316)', 'Chunking + embedding, FTS5 + ChromaDB indexing, PageIndex tree building'],
            ['Week 3 (Mar 17\u201323)', 'Four retrieval methods + RRF fusion, Ollama integration, RAG pipeline'],
            ['Week 4 (Mar 24\u201330)', 'Streamlit UI with source tracing, evaluation (20 test questions)'],
            ['Week 5 (Mar 31\u2013Apr 3)', 'ROS 2 integration (Part 2), final report, presentation'],
        ],
        col_widths=[4.5, 13.0]
    )

    p = doc.add_paragraph()
    r = p.add_run('Evaluation Method: ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        'We will prepare 20 domain-specific questions spanning multiple textbooks, '
        'run the system on each, and manually assess correctness using a 3-level '
        'scale (1 = correct, 0.5 = partially correct, 0 = incorrect). For each '
        'question, we will also list the top 3 retrieved chunks and mark their '
        'relevance. The final accuracy score will be the average across all 20 questions.'
    ).font.size = Pt(11)

    # ── EXPECTED RESULTS ──
    doc.add_heading('Expected Results', level=1)
    results = [
        'A functional RAG-based Q&A system that answers AI/ML educational questions with >80% accuracy',
        'Deep source tracing with clickable references that highlight the exact region on original PDF pages',
        'Layout-aware processing that correctly handles tables, formulas, and figures without losing structure',
        'A responsive Streamlit UI for interactive querying',
    ]
    for r_text in results:
        p = doc.add_paragraph(r_text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)

    # ── WHY IS IT INTERESTING ──
    doc.add_heading('Why Is It Interesting?', level=1)
    interests = [
        ('Deep source tracing', 'Goes beyond basic RAG by providing pixel-level traceability to original documents, not just page references'),
        ('Layout-aware parsing', 'Uses YOLO-based document layout detection to preserve tables, formulas, and figures that traditional text extraction would break'),
        ('Practical utility', 'Directly useful for students studying AI/ML across multiple textbooks'),
        ('Scalable', 'The same architecture can be applied to any domain of knowledge'),
        ('Low resource', 'Runs entirely locally with <1.5GB model, suitable for deployment on constrained hardware (robots)'),
    ]
    for i, (label, desc) in enumerate(interests, 1):
        p = doc.add_paragraph(f'{i}. ')
        r = p.add_run(f'{label} — ')
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assignment2_proposal.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


if __name__ == '__main__':
    build()
